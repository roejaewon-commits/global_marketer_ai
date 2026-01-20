import base64
import datetime as dt
import pandas as pd
import requests
import streamlit as st
import fitz  # PyMuPDF
import urllib3
import io
from openai import OpenAI
from duckduckgo_search import DDGS
from pytrends.request import TrendReq
from docx import Document
from docx.shared import Pt

# ---------------------------------------------------------
# 0. 초기 설정
# ---------------------------------------------------------
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
st.set_page_config(page_title="AI 글로벌 마케터 (V11.1)", layout="wide")

def get_secret(key: str) -> str:
    val = st.secrets.get(key, "")
    return val.strip() if val else ""

OPENAI_API_KEY = get_secret("OPENAI_API_KEY")
KOTRA_SERVICE_KEY = get_secret("KOTRA_SERVICE_KEY")

# ---------------------------------------------------------
# 1. 스마트 국가코드 변환기
# ---------------------------------------------------------
def get_smart_country_code(user_input):
    mapping = {
        "대한민국": "KR", "한국": "KR", "KOREA": "KR", "SOUTH KOREA": "KR",
        "미국": "US", "USA": "US", "AMERICA": "US",
        "중국": "CN", "CHINA": "CN",
        "일본": "JP", "JAPAN": "JP",
        "베트남": "VN", "VIETNAM": "VN",
        "인도네시아": "ID", "INDONESIA": "ID", "인니": "ID", "INA": "ID",
        "태국": "TH", "THAILAND": "TH",
        "인도": "IN", "INDIA": "IN",
        "독일": "DE", "GERMANY": "DE",
        "프랑스": "FR", "FRANCE": "FR",
        "영국": "GB", "UK": "GB",
        "호주": "AU", "AUSTRALIA": "AU"
    }
    clean_input = user_input.upper().strip()
    if clean_input in mapping: return mapping[clean_input]
    if len(clean_input) == 2: return clean_input
    return None

# ---------------------------------------------------------
# 2. 세션 및 Word 생성기
# ---------------------------------------------------------
if "inputs" not in st.session_state:
    st.session_state.inputs = {
        "company_name": "숭실시스템즈",
        "country_input": "인도네시아", 
        "real_code": "ID",           
        "keyword": "Food Packaging", 
        "budget": 5000000
    }

if "vision_analysis" not in st.session_state: st.session_state.vision_analysis = ""
if "market_data" not in st.session_state: st.session_state.market_data = {"macro": {}, "report": "", "trends": pd.DataFrame()}
if "final_report" not in st.session_state: st.session_state.final_report = ""
if "emails" not in st.session_state: st.session_state.emails = {"KR": "", "EN": ""}
if "sns_content" not in st.session_state: st.session_state.sns_content = {"Insta_KR": "", "Insta_EN": "", "Linked_KR": "", "Linked_EN": ""}

def create_word_docx(company, country, vision, report, emails):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Malgun Gothic'
    style.font.size = Pt(10)

    doc.add_heading(f'{company} - {country} 진출 전략 보고서', 0)
    doc.add_paragraph(f"생성 일자: {dt.date.today()}")
    
    doc.add_heading('1. 제품 및 내부 역량 정밀 분석', level=1)
    doc.add_paragraph(vision)
    
    doc.add_heading('2. 시장 진입 전략', level=1)
    doc.add_paragraph(report)
    
    doc.add_heading('3. B2B 영업 제안 메일', level=1)
    doc.add_heading('[국문]', level=2)
    doc.add_paragraph(emails.get("KR", ""))
    doc.add_heading('[English]', level=2)
    doc.add_paragraph(emails.get("EN", ""))
    
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# ---------------------------------------------------------
# 3. 분석 및 생성 모듈 (Vision 프롬프트 강화됨!)
# ---------------------------------------------------------
def analyze_pdf_with_vision(uploaded_file):
    if not OPENAI_API_KEY: return "API Key 필요"
    doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
    max_pages = 3 
    base64_images = []
    for i in range(min(len(doc), max_pages)):
        page = doc.load_page(i)
        pix = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5))
        img_data = pix.tobytes("png")
        base64_images.append(base64.b64encode(img_data).decode('utf-8'))
    
    client = OpenAI(api_key=OPENAI_API_KEY)
    
    # [수정됨] 디테일 강화 프롬프트
    prompt = """
    당신은 20년 경력의 수석 기술 마케터입니다. 업로드된 카탈로그(PDF)를 정밀 분석하여 보고서를 작성하세요.
    단순한 요약이 아니라, 카탈로그에 있는 **구체적인 스펙, 수치, 인증 마크, 기술 용어**를 인용하여 전문성 있게 작성해야 합니다.
    
    [분석 항목]
    1. **핵심 제품 포트폴리오 (Core Products)**:
       - 주요 제품 라인업을 나열하고 각각의 특징을 구체적으로 설명하세요.
    2. **기술적 차별점 (Technical USP)**:
       - 경쟁사 대비 돋보이는 기술, 특허, 정밀도, 속도, 소재(SUS 등) 등의 스펙을 찾아내어 강조하세요.
       - HACCP, GMP 등 인증 마크가 보이면 반드시 언급하세요.
    3. **고객 도입 효과 (Customer Benefits)**:
       - 이 기계를 도입했을 때 공장이 얻게 되는 이득(생산성 향상, 이물질 사고 예방 등)을 구체적으로 서술하세요.
    4. **추천 타겟 산업**:
       - 이 제품이 가장 필요한 산업군(예: 제과, 육가공, 수산물 등)을 추론하세요.
       
    [작성 지침]
    - 각 항목당 최소 3~5문장으로 상세하게 작성하세요.
    - 톤앤매너: 신뢰감 있고 전문적인 비즈니스 어조.
    """
    
    payload = [{"type": "text", "text": prompt}]
    for b64 in base64_images:
        payload.append({"type": "image_url", "image_url": {"url": f"data:image/png;base64,{b64}"}})
    
    res = client.chat.completions.create(model="gpt-4o", messages=[{"role": "user", "content": payload}])
    return res.choices[0].message.content

def fetch_rich_macro_economics(country_code):
    indicators = {
        "NY.GDP.MKTP.CD": "GDP (시장규모)", "NY.GDP.MKTP.KD.ZG": "경제성장률",
        "SP.POP.TOTL": "총 인구수", "NY.GNP.PCAP.CD": "1인당 GNI",
        "FP.CPI.TOTL.ZG": "물가상승률", "IT.NET.USER.ZS": "인터넷 사용률"
    }
    macro_data = {}
    for code, name in indicators.items():
        try:
            url = f"http://api.worldbank.org/v2/country/{country_code}/indicator/{code}?format=json&per_page=1&date=2021:2024"
            r = requests.get(url, timeout=3)
            if r.status_code == 200:
                data = r.json()
                if len(data) > 1 and data[1]:
                    item = data[1][0]
                    val = item['value']
                    year = item['date']
                    if val is None: disp = "N/A"
                    else:
                        if "GDP" in name: disp = f"${val/1e9:,.1f} B"
                        elif "인구수" in name: disp = f"{val/1e6:,.1f} M"
                        elif "GNI" in name: disp = f"${val:,.0f}"
                        else: disp = f"{val:.1f}%"
                    macro_data[name] = {"value": disp, "year": year}
                else: macro_data[name] = {"value": "N/A", "year": "-"}
        except: macro_data[name] = {"value": "N/A", "year": "-"}
    return macro_data

def fetch_industry_report(country, keyword):
    client = OpenAI(api_key=OPENAI_API_KEY)
    queries = [f"{country} {keyword} market size 2025", f"{country} {keyword} trends", f"top {keyword} companies in {country}"]
    txt = ""
    with DDGS() as ddgs:
        for q in queries:
            try:
                for r in list(ddgs.text(q, max_results=2)): txt += f"- {r['title']}: {r['body']}\n"
            except: pass
    if not txt: return "정보 부족"
    
    prompt = f"'{country} {keyword} 시장 리포트' 작성. 기준 연도 명시. [정보] {txt}"
    res = client.chat.completions.create(model="gpt-4o", messages=[{"role": "user", "content": prompt}])
    return res.choices[0].message.content

def fetch_all_intelligence(inputs):
    macro = fetch_rich_macro_economics(inputs['real_code'])
    report = fetch_industry_report(inputs['country_input'], inputs['keyword'])
    trend_df = pd.DataFrame()
    try:
        pytrend = TrendReq(hl='en-US', tz=360, timeout=(5,5))
        pytrend.build_payload([inputs['keyword']], cat=0, timeframe='today 12-m', geo=inputs['real_code'])
        trend_df = pytrend.interest_over_time()
    except: pass
    return {"macro": macro, "report": report, "trends": trend_df}

def generate_strategy(inputs, vision, mkt_data):
    client = OpenAI(api_key=OPENAI_API_KEY)
    macro_info = "\n".join([f"{k}: {v['value']} ({v['year']})" for k, v in mkt_data['macro'].items()])
    prompt = f"전략보고서 작성. 기업:{inputs['company_name']}->{inputs['country_input']}. 예산:{inputs['budget']:,}원. \n[제품]{vision}\n[시장]{macro_info}\n[트렌드]{mkt_data['report']}"
    res = client.chat.completions.create(model="gpt-4o", messages=[{"role": "user", "content": prompt}])
    return res.choices[0].message.content

def generate_email(inputs, vision, lang):
    client = OpenAI(api_key=OPENAI_API_KEY)
    prompt = f"B2B 영업메일 작성. 언어:{lang}. 타겟:{inputs['country_input']}. 제품:{vision}"
    res = client.chat.completions.create(model="gpt-4o", messages=[{"role": "user", "content": prompt}])
    return res.choices[0].message.content

def generate_sns(inputs, vision, plat, lang):
    client = OpenAI(api_key=OPENAI_API_KEY)
    lang_instruction = "MUST be written in KOREAN." if lang == "Korean" else "MUST be written in ENGLISH."
    style = "감성적이고 트렌디한 인스타그램 스타일 (해시태그 포함)" if plat == "Instagram" else "전문적인 링크드인 비즈니스 스타일"
    
    prompt = f"""
    Create a {plat} post for {inputs['company_name']}.
    Target Market: {inputs['country_input']}
    Product Info: {vision}
    Style: {style}
    IMPORTANT: The output language {lang_instruction}
    """
    res = client.chat.completions.create(model="gpt-4o", messages=[{"role": "user", "content": prompt}])
    return res.choices[0].message.content

# ---------------------------------------------------------
# 4. 메인 UI
# ---------------------------------------------------------
st.title("🌏 AI 글로벌 마케터 (V11.1)")
st.caption("Vision 분석 디테일 강화 + SNS 다국어 + 키워드 최적화")

with st.sidebar:
    st.header("⚙️ 설정")
    st.session_state.inputs["company_name"] = st.text_input("회사명", st.session_state.inputs["company_name"])
    
    user_country = st.text_input("진출 국가 (한글/영어)", st.session_state.inputs["country_input"])
    detected_code = get_smart_country_code(user_country)
    
    if detected_code:
        st.session_state.inputs["country_input"] = user_country
        st.session_state.inputs["real_code"] = detected_code
        st.success(f"✅ 감지됨: {user_country} ({detected_code})")
    else:
        st.error("⚠️ 국가 식별 불가")
        st.session_state.inputs["real_code"] = ""

    st.session_state.inputs["keyword"] = st.text_input("트렌드 키워드 (영어)", st.session_state.inputs["keyword"], help="구글 트렌드 검색용")
    
    budget_val = st.number_input("마케팅 예산", value=st.session_state.inputs["budget"], step=1000000)
    st.session_state.inputs["budget"] = budget_val
    st.caption(f"💰 {budget_val:,} 원")
    
    if st.button("🔄 리셋"): st.session_state.clear(); st.rerun()

tabs = st.tabs(["1️⃣ 제품 분석", "2️⃣ 시장 인텔리전스", "3️⃣ 전략 보고서", "4️⃣ 영업 메일", "5️⃣ SNS 콘텐츠", "📥 다운로드"])

with tabs[0]:
    st.subheader("👁️ Vision 제품 분석 (Deep Analysis)")
    f = st.file_uploader("PDF 업로드", type="pdf")
    if f and st.button("정밀 분석 시작"):
        with st.spinner("AI가 카탈로그를 정밀 분석 중입니다 (시간이 조금 더 소요될 수 있습니다)..."):
            st.session_state.vision_analysis = analyze_pdf_with_vision(f)
            st.success("분석 완료")
    if st.session_state.vision_analysis: st.info(st.session_state.vision_analysis)

with tabs[1]:
    st.subheader("📊 국가 & 산업 심층 분석")
    if st.button("데이터 분석 실행"):
        with st.spinner(f"{st.session_state.inputs['country_input']} 시장 분석 중..."):
            st.session_state.market_data = fetch_all_intelligence(st.session_state.inputs)
            st.success("완료")
    
    data = st.session_state.market_data
    macro = data.get("macro", {})
    if macro:
        st.markdown(f"### 🚩 {st.session_state.inputs['country_input']} 핵심 지표 Dashboard")
        c1, c2, c3 = st.columns(3)
        c1.metric("GDP", macro.get("GDP (시장규모)", {}).get("value", "-"), help=f"기준: {macro.get('GDP (시장규모)', {}).get('year')}")
        c2.metric("경제성장률", macro.get("경제성장률", {}).get("value", "-"))
        c3.metric("인구수", macro.get("총 인구수", {}).get("value", "-"))
        c4, c5, c6 = st.columns(3)
        c4.metric("1인당 GNI", macro.get("1인당 GNI", {}).get("value", "-"))
        c5.metric("물가상승률", macro.get("물가상승률", {}).get("value", "-"))
        c6.metric("인터넷 사용률", macro.get("인터넷 사용률", {}).get("value", "-"))
        st.divider()

    if data['report']:
        st.markdown(f"### 📑 {st.session_state.inputs['keyword']} 산업 리포트")
        st.write(data['report'])
    
    if not data['trends'].empty:
        st.line_chart(data['trends'])

with tabs[2]:
    st.subheader("📑 전략 보고서")
    if st.button("보고서 생성"):
        with st.spinner("작성 중..."):
            st.session_state.final_report = generate_strategy(st.session_state.inputs, st.session_state.vision_analysis, st.session_state.market_data)
    if st.session_state.final_report: st.markdown(st.session_state.final_report)

with tabs[3]:
    st.subheader("✉️ 영업 메일")
    if st.button("메일 생성"):
        with st.spinner("작성 중..."):
            st.session_state.emails["KR"] = generate_email(st.session_state.inputs, st.session_state.vision_analysis, "Korean")
            st.session_state.emails["EN"] = generate_email(st.session_state.inputs, st.session_state.vision_analysis, "English")
    if st.session_state.emails["KR"]:
        t1, t2 = st.tabs(["KR", "EN"])
        with t1: st.text_area("Korean", st.session_state.emails["KR"], height=400)
        with t2: st.text_area("English", st.session_state.emails["EN"], height=400)

with tabs[4]:
    st.subheader("📱 SNS 콘텐츠 (다국어 지원)")
    if st.button("콘텐츠 생성 (4종)"):
        with st.spinner("인스타그램 및 링크드인 게시물 생성 중..."):
            st.session_state.sns_content["Insta_KR"] = generate_sns(st.session_state.inputs, st.session_state.vision_analysis, "Instagram", "Korean")
            st.session_state.sns_content["Insta_EN"] = generate_sns(st.session_state.inputs, st.session_state.vision_analysis, "Instagram", "English")
            st.session_state.sns_content["Linked_KR"] = generate_sns(st.session_state.inputs, st.session_state.vision_analysis, "LinkedIn", "Korean")
            st.session_state.sns_content["Linked_EN"] = generate_sns(st.session_state.inputs, st.session_state.vision_analysis, "LinkedIn", "English")
            st.success("완료!")
            
    if st.session_state.sns_content["Insta_KR"]:
        s1, s2 = st.tabs(["📸 Instagram", "💼 LinkedIn"])
        with s1:
            c1, c2 = st.columns(2)
            with c1: st.text_area("Insta KR", st.session_state.sns_content["Insta_KR"], height=400)
            with c2: st.text_area("Insta EN", st.session_state.sns_content["Insta_EN"], height=400)
        with s2:
            c1, c2 = st.columns(2)
            with c1: st.text_area("Linked KR", st.session_state.sns_content["Linked_KR"], height=400)
            with c2: st.text_area("Linked EN", st.session_state.sns_content["Linked_EN"], height=400)

with tabs[5]:
    st.subheader("📥 결과물 다운로드")
    ready = st.session_state.final_report and st.session_state.emails["KR"]
    if ready:
        docx = create_word_docx(
            st.session_state.inputs["company_name"],
            st.session_state.inputs["country_input"],
            st.session_state.vision_analysis,
            st.session_state.final_report,
            st.session_state.emails
        )
        st.download_button(
            label="📄 Word 보고서 다운로드",
            data=docx,
            file_name=f"Strategy_{st.session_state.inputs['company_name']}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    else:
        st.info("먼저 [3.전략 보고서]와 [4.영업 메일]을 생성해주세요.")
